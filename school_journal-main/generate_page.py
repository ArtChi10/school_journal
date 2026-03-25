import os
import marks_dict as m

from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
import helpers as h
from helpers import get_assistant_principal_for_class

def _fmt_num(n):
    """25.0 -> '25', 25.50 -> '25.5', '25.0' -> '25'"""
    if n is None or n == "":
        return ""
    try:
        f = float(n)
        return str(int(f)) if f.is_integer() else ("{:.10g}".format(f)).rstrip(".")
    except (TypeError, ValueError):
        return str(n)

# Function to set cell background color
def set_cell_background_color(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def _no_mark(v):
    """
    True, если по критерию нет оценки (учитель не выбрал ни одной).
    У тебя такие случаи приходят как [-1, 'NOT_FOUND'].
    """
    if v is None:
        return True
    if isinstance(v, (list, tuple)):
        # пустая/битая структура
        if len(v) < 2:
            return True
        # твой реальный маркер пустоты:
        if v[0] == -1 or (isinstance(v[1], str) and v[1] == 'NOT_FOUND'):
            return True
    # строки типа "" и пр.
    if isinstance(v, str) and v.strip() == "":
        return True
    return False

LETTER_TO_LABEL = {
    "A*": "Превосходит ожидания по программе | Exceeds expectations (90-100%)",
    "A+": "Превосходит ожидания по программе | Exceeds expectations (90-100%)",
    "A": "Соответствует ожиданиям по программе | Meets expectations (50-89%)",
    "B": "Соответствует ожиданиям по программе | Meets expectations (50-89%)",
    "C": "Соответствует ожиданиям по программе | Meets expectations (50-89%)",
    "D": "Соответствует ожиданиям по программе | Meets expectations (50-89%)",
    "E": "Начальный уровень | Beginner level (40-49%)",
    "F": "Не проявляется | Not evident (0-39%)",
    "G": "Не проявляется | Not evident (0-39%)",
    "U": "Не проявляется | Not evident (0-39%)",
}
LETTER_SET = set(LETTER_TO_LABEL.keys())

def _is_numberlike(x):
    try:
        float(x)
        return True
    except (TypeError, ValueError):
        return False

def _bucket_by_percent(pct, marks_dict):
    """
    Разложить процент по полуинтервалам [lo, hi) из marks_dict.
    -1 трактуем как «нет оценки».
    """
    if pct is None:
        return -1, "NOT_FOUND"
    try:
        p = float(pct)
    except Exception:
        return -1, "NOT_FOUND"

    for label, meta in marks_dict.items():
        lo, hi = meta["bounds"]
        if lo == -1:
            if p < hi:   # (-∞, hi)
                return meta["order"], label
        else:
            if lo <= p < hi:  # [lo, hi)
                return meta["order"], label
    return -1, "NOT_FOUND"

def make_mark_value(cell_value, marks_dict):
    """
    Универсальный парсер значения из XLSX:
    - 'A'..'F'  -> [order, label, 'A'..'F']
    - число     -> [order, label, <процент>]
    - '-', ''   -> [-1, 'NOT_FOUND']
    """
    if cell_value is None:
        return [-1, "NOT_FOUND"]

    s = str(cell_value).strip()
    if s == "" or s == "-":
        return [-1, "NOT_FOUND"]

    # Буква A..F
    up = s.upper()
    if up in LETTER_SET:
        label = LETTER_TO_LABEL[up]
        order = marks_dict[label]["order"]
        return [order, label, up]

    # Число (проценты)
    if _is_numberlike(s):
        order, label = _bucket_by_percent(float(s), marks_dict)
        if order == -1:
            return [-1, "NOT_FOUND"]
        return [order, label, float(s)]

    # Иначе не распознали
    return [-1, "NOT_FOUND"]


def _is_numberlike(x):
    try:
        float(x)
        return True
    except (TypeError, ValueError):
        return False

def _extract_letter(value):
    """
    Ищет букву A..F:
    - если value = 'A'/'B'/... (строка) — вернёт букву;
    - если value = [.., .., 'A'] — возьмёт третий элемент как букву;
    - иначе None.
    """
    if isinstance(value, str):
        s = value.strip().upper()
        return s if s in LETTER_SET else None
    if isinstance(value, (list, tuple)) and len(value) >= 3 and isinstance(value[2], str):
        s = value[2].strip().upper()
        return s if s in LETTER_SET else None
    return None



def _add_full_width_row(table):
    """
    Добавляет строку в конец таблицы и склеивает все ячейки в одну.
    Возвращает эту единую (merged) ячейку.
    """
    row = table.add_row()
    merged = row.cells[0]
    for c in range(1, len(row.cells)):
        merged = merged.merge(row.cells[c])
    merged.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = merged.paragraphs[0]
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    return merged


# Function to set cell border color
def set_cell_border_color(cell, color):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    for border_name in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcPr.append(border)


# Function to set table width
def set_table_width(table, width_cm):
    table_element = table._tbl
    tbl_pr = table_element.tblPr
    tbl_w = OxmlElement('w:tblW')
    tbl_w.set(qn('w:w'), str(int(width_cm) * 597))  # Word width unit is 1/20 pt, 1 cm = 567 twips
    tbl_w.set(qn('w:type'), 'dxa')
    tbl_pr.append(tbl_w)


def generate_subject(students, output_path):
    paths = []
    # for fixed "Критерии оценивания" column
    max_cols = len(m.marks_dict) + 1
    for s in students:
        # print(f'Processing {s.last_name} {s.first_name}')
        # ==== DEBUG: посмотреть все значения s.subjects ====
        # ================================================
        for sub in s.subjects:
            # print(f'Subject: {sub.name}')
            criteria_count = len(sub.marks)
            # +4 rows for static rows
            max_rows = criteria_count + 4

            doc = Document()

            # Set default font to Arial
            styles = doc.styles
            font = styles['Normal'].font
            font.name = 'Arial'
            font.size = Pt(12)

            # Create the table with specified structure
            table = doc.add_table(rows=max_rows, cols=max_cols)
            from docx.shared import Cm
            from docx.oxml import OxmlElement
            from docx.oxml.ns import qn

            def _cm_to_twips(c):  # 1 inch = 1440 twips; 1 inch = 2.54 cm
                return int(c / 2.54 * 1440)

            def set_col_widths_cm(table, widths_cm):
                """
                Жёстко задаёт ширины колонок:
                - прописывает tblGrid (w:gridCol w:w=...),
                - проставляет w:tcW type=dxa каждой ячейке,
                - отключает авто-подгонку Word.
                """
                table.autofit = False

                # 1) tblGrid
                tbl = table._tbl
                # удалить старую сетку, если есть
                if tbl.tblGrid is not None:
                    tbl.remove(tbl.tblGrid)
                grid = OxmlElement('w:tblGrid')
                tbl.insert(0, grid)
                for wcm in widths_cm:
                    gcol = OxmlElement('w:gridCol')
                    gcol.set(qn('w:w'), str(_cm_to_twips(wcm)))
                    grid.append(gcol)

                # 2) tcW для каждой ячейки + .width как подсказка
                for row in table.rows:
                    for j, wcm in enumerate(widths_cm):
                        cell = row.cells[j]
                        cell.width = Cm(wcm)
                        tcPr = cell._tc.get_or_add_tcPr()
                        # удалить старый tcW, если присутствует
                        for el in list(tcPr):
                            if el.tag == qn('w:tcW'):
                                tcPr.remove(el)
                        tcW = OxmlElement('w:tcW')
                        tcW.set(qn('w:w'), str(_cm_to_twips(wcm)))
                        tcW.set(qn('w:type'), 'dxa')
                        tcPr.append(tcW)

            # -------- вызов: первую колонку делаем широкой --------
            section = doc.sections[0]

            # EMU → cm: 1 см = 360000 EMU
            EMU_PER_CM = 360000
            usable_width_emu = section.page_width - section.left_margin - section.right_margin
            usable_width_cm = usable_width_emu / EMU_PER_CM  # float, в сантиметрах

            first_cm = 6.5  # ПОИГРАЙ ЗНАЧЕНИЕМ: 8.5–10.0 см даст заметную ширину
            rest_cm = (usable_width_cm - first_cm) / (max_cols - 1)
            widths = [first_cm] + [rest_cm] * (max_cols - 1)

            set_col_widths_cm(table, widths)
            # -------- конец блока --------

            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            table.style = 'TableGrid'

            set_table_width(table, 18)
            # Set column spans for the first and second rows
            table.cell(0, 0).merge(table.cell(0, max_cols - 1))
            table.cell(1, 0).merge(table.cell(1, max_cols - 1))

            # Fill in the content and format the Subject cell
            subject_cell = table.cell(0, 0)
            p = subject_cell.paragraphs[0]
            run = p.add_run(f'\n{sub.name.title()}\n')
            run.bold = True
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            subject_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background_color(subject_cell, '5885e1')

            # Description
            descriptor_cell = table.cell(1, 0)
            p = descriptor_cell.paragraphs[0]
            run = p.add_run('Дескриптор:\n')
            run.bold = True
            run = p.add_run(sub.descriptor)

            # Criteria columns
            # --- Criteria columns (заголовок строки критериев + названия столбцов оценок) ---
            criteria_cell = table.cell(2, 0)
            criteria_cell.text = ""  # очищаем, чтобы управлять run'ом
            p = criteria_cell.paragraphs[0]
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            r = p.add_run("Критерии оценивания | Assessment criteria")
            r.bold = True
            criteria_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_background_color(criteria_cell, 'ccdaf5')

            # заголовки оценочных колонок — меньший кегль и по центру
            for j, mark in enumerate(m.marks_dict, start=1):
                hdr = table.cell(2, j)
                hdr.text = ""  # очищаем
                set_cell_background_color(hdr, 'ccdaf5')
                hdr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = hdr.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                r = p.add_run(mark)
                r.font.size = Pt(7)  # <-- размер 9
            # --- /Criteria columns ---


            # Criteria rows
            i = 3
            for key, value in sub.marks.items():
                # заголовок критерия
                cell = table.cell(i, 0)
                cell.text = "" if key is None else str(key)

                j = 1
                #print(type(value), value[0])
                if _no_mark(value):
                    # НЕТ оценки → ставим длинные тире во всех колонках
                    for mark in m.marks_dict:
                        cell = table.cell(i, j)
                        cell.text = '—'
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        j += 1
                elif value[0] == -2:
                    # value = [-2, "57 B"] или "B 57" / "57% B" и т.п.
                    raw = str(value[1]).strip().replace("%", "")
                    num_str = None
                    letter = None

                    # достаём число и букву (A, B, C, D, E, F, A*)
                    for tok in raw.split():
                        up = tok.upper()
                        if up in LETTER_SET:
                            letter = up
                        else:
                            try:
                                num = float(tok.replace(",", "."))
                                # красивое представление без лишних .0
                                num_str = str(int(num)) if num.is_integer() else str(num)
                            except Exception:
                                pass
                    if num_str is None:
                        display = f"{letter}".strip() if letter else (num_str or raw)
                    else:
                        display = f"{num_str}% {letter}".strip() if letter else (num_str or raw)

                    # целевая колонка по букве
                    target_label = LETTER_TO_LABEL.get(letter) if letter else None

                    for mark in m.marks_dict:
                        cell = table.cell(i, j)
                        if target_label and mark == target_label:
                            cell.text = display  # <- пишем ровно "57 B"
                        else:
                            cell.text = ""  # остальные столбцы — пусто
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for p in cell.paragraphs:
                            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        j += 1

                else:
                    # оценка есть → как раньше: отмечаем нужную колонку
                    for mark in m.marks_dict:
                        cell = table.cell(i, j)
                        if value[1] == mark:
                            if len(value) == 3:
                                #cell.text = f"{value[2]}%"
                                cell.text = f"{_fmt_num(value[2])}%"
                            else:
                                cell.text = 'V'
                        else:
                            #print(value[1])
                            cell.text = ''
                        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                        for paragraph in cell.paragraphs:
                            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        j += 1
                i += 1
            # --- Комментарий по предмету (если есть) ---
            has_comment = bool(getattr(sub, "comment", None))
            if has_comment:
                comment_row_idx = max_rows - 1
                cell = table.cell(comment_row_idx, 0)
                cell = cell.merge(table.cell(comment_row_idx, max_cols - 1))
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                p = cell.paragraphs[0]
                p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                run = p.add_run("Примечание | Comment: ")
                run.bold = True
                p.add_run(str(sub.comment).strip())

                # HIGHLIGHT IF RETAKE = "да"
                # допускаем разные имена атрибута: retake / resit
                retake_raw = getattr(sub, "retake", getattr(sub, "resit", ""))
                need_highlight = str(retake_raw).strip().lower() in {"да", "yes", "y", "true", "1"}
                if need_highlight:
                    from docx.oxml import OxmlElement
                    from docx.oxml.ns import qn
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement("w:shd")
                    shd.set(qn("w:val"), "clear")
                    shd.set(qn("w:color"), "auto")
                    shd.set(qn("w:fill"), "FFFF00")  # жёлтый
                    tcPr.append(shd)
            # --- /Комментарий ---

            # --- Учитель ---
            if has_comment:
                cell = _add_full_width_row(table)
            else:
                teacher_row_idx = max_rows - 1
                cell = table.cell(teacher_row_idx, 0)
                cell = cell.merge(table.cell(teacher_row_idx, max_cols - 1))

            p = cell.paragraphs[0]
            run = p.add_run("Учитель | Teacher: ")
            run.bold = True
            p.add_run(str(sub.teacher).strip())
            # --- /Учитель ---


            # Set border color for all cells
            for row in table.rows:
                for cell in row.cells:
                    set_cell_border_color(cell, '000000')

            # Save the document
            to_print = {'StudentName': f'{s.first_name} {s.last_name}', 'Class': s.level}
            fname = h.sanitize_filename(f"{sub.name.title()} {to_print['StudentName']}.docx")
            h.ensure_dir(output_path)
            filepath = os.path.join(output_path, fname)
            doc.save(filepath)
            paths.append({'student': f"{to_print['StudentName']}{to_print['Class']}",
                          'filepath': filepath,
                          'filename': f"{to_print['StudentName']}.docx"})
    return paths